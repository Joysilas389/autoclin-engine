"""
AutoClin Engine DOCX Narrative Report Generator
Plain-language Word document for clinical study reports / regulatory submissions.
"""
from datetime import datetime, timezone
from collections import Counter


class DOCXReportGenerator:
    """Generates a narrative Word document from pipeline results."""

    def generate(self, result, output_path: str, dataset_name: str = "Dataset"):
        try:
            from docx import Document
            from docx.shared import Inches, Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            raise RuntimeError("python-docx required: pip install python-docx")

        doc = Document()

        # Styles
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15

        # ── Title Page ──
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("AutoClin Engine\n")
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(30, 64, 175)
        run.bold = True
        run = p.add_run("Data Quality Assessment Report\n\n")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(71, 85, 105)

        now = datetime.now(timezone.utc).strftime("%B %d, %Y")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Dataset: {dataset_name}\n").font.size = Pt(12)
        p.add_run(f"Report Date: {now}\n").font.size = Pt(11)
        p.add_run(f"AutoClin Engine Version: 1.0.0\n").font.size = Pt(11)
        p.add_run(f"\nCONFIDENTIAL — FOR INTERNAL USE ONLY").font.size = Pt(10)
        doc.add_page_break()

        # ── 1. Introduction ──
        doc.add_heading("1. Introduction", level=1)
        doc.add_paragraph(
            f"This report presents the results of an automated data quality assessment "
            f"performed by AutoClin Engine on the dataset '{dataset_name}'. The assessment "
            f"employed {len(result.method_rankings)} unsupervised machine learning methods "
            f"to detect anomalies, noise, and data quality issues without requiring "
            f"labeled training data or manual rules."
        )
        doc.add_paragraph(
            f"The dataset contains {result.total_rows:,} records. AutoClin Engine's "
            f"pipeline ran through 9 automated phases: ingestion, profiling, clinical "
            f"field mapping, preprocessing, multi-engine detection, method benchmarking, "
            f"ensemble selection, cleaning recommendation, and reporting."
        )

        # ── 2. Methods ──
        doc.add_heading("2. Methods", level=1)
        doc.add_paragraph(
            f"AutoClin Engine applied the following unsupervised detection methods: "
            f"{', '.join(m['method'] for m in result.method_rankings)}. "
            f"Each method was scored on six dimensions: Anomaly Discrimination (AD), "
            f"Stability Score (SS), Clinical Plausibility (CP), Explainability (EX), "
            f"Computational Cost (CC), and Noise Detection Confidence (NDC)."
        )
        doc.add_paragraph(
            f"The method selection mode was '{result.selection_mode}'. "
            f"The selected method(s): {', '.join(result.selected_methods)}."
        )

        # Method table
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Method"
        hdr[1].text = "Composite"
        hdr[2].text = "NDC"
        hdr[3].text = "Selected"
        for mr in result.method_rankings:
            row = table.add_row().cells
            row[0].text = mr["method"]
            row[1].text = f"{mr['composite']:.4f}"
            row[2].text = f"{mr['ndc']:.4f}"
            row[3].text = "Yes" if mr.get("selected") else ""

        # ── 3. Results ──
        doc.add_heading("3. Results", level=1)

        bm = result.before_metrics
        doc.add_heading("3.1 Key Findings", level=2)
        doc.add_paragraph(
            f"AutoClin Engine detected {result.total_anomalies} anomalies across "
            f"{result.total_rows:,} records, yielding a noise rate of "
            f"{result.noise_percentage:.2f}%. The overall Data Trust Score is "
            f"{result.trust_score:.1f} out of 100."
        )

        # Metrics table
        doc.add_heading("3.2 Quality Metrics", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        table.rows[0].cells[0].text = "Metric"
        table.rows[0].cells[1].text = "Value"
        metrics = [
            ("Noise Rate", f"{bm.get('noise_pct', 0):.2f}%"),
            ("Trust Score", f"{bm.get('trust_score', 0):.1f} / 100"),
            ("Missingness Rate", f"{bm.get('missingness_pct', 0):.2f}%"),
            ("Duplicate Rate", f"{bm.get('duplicate_pct', 0):.2f}%"),
            ("Clinical Plausibility", f"{bm.get('plausibility_rate', 0):.1f}%"),
        ]
        for label, value in metrics:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = value

        # Anomaly summary
        doc.add_heading("3.3 Anomaly Summary", level=2)
        type_dist = Counter(c.anomaly_type for c in result.anomaly_classifications)
        sev_dist = Counter(c.severity for c in result.anomaly_classifications)

        doc.add_paragraph(result.global_summary.get("narrative", ""))

        doc.add_paragraph(
            f"Severity distribution: "
            f"{sev_dist.get('critical', 0)} critical, "
            f"{sev_dist.get('high', 0)} high, "
            f"{sev_dist.get('medium', 0)} medium, "
            f"{sev_dist.get('low', 0)} low."
        )

        # ── 4. Discussion ──
        doc.add_heading("4. Discussion", level=1)
        doc.add_paragraph(
            "The automated assessment identified several data quality issues that "
            "warrant manual review. The following limitations should be considered:"
        )
        doc.add_paragraph(
            "• Unsupervised methods may produce false positives, particularly for "
            "rare but clinically valid observations.",
            style="List Bullet",
        )
        doc.add_paragraph(
            "• Clinical plausibility checks rely on reference ranges that may not "
            "account for all clinical contexts.",
            style="List Bullet",
        )
        doc.add_paragraph(
            "• The selected method's performance may vary with different data "
            "characteristics. The stability score provides an indication of "
            "reproducibility.",
            style="List Bullet",
        )

        # ── 5. Recommendations ──
        doc.add_heading("5. Recommendations", level=1)
        action_dist = Counter(r.recommended_action for r in result.cleaning_recommendations)
        doc.add_paragraph(
            f"AutoClin Engine recommends {len(result.cleaning_recommendations)} cleaning actions. "
            f"The most common recommended action is "
            f"'{max(action_dist, key=action_dist.get)}' "
            f"({action_dist[max(action_dist, key=action_dist.get)]} instances)."
        )
        doc.add_paragraph(
            "All flagged records should be reviewed by a qualified clinical data manager "
            "or clinician before corrections are applied. Auto-cleaning should only be "
            "used for deterministic, reversible corrections such as unit conversions and "
            "encoding normalization."
        )

        # ── Disclaimer ──
        doc.add_heading("Disclaimer", level=1)
        doc.add_paragraph(
            "This report was generated by AutoClin Engine, an automated unsupervised ML tool. "
            "Findings should NOT be interpreted as definitive clinical judgments. All "
            "recommendations must be reviewed by qualified professionals before action. "
            "AutoClin Engine supports but does not replace regulatory-compliant data management."
        )

        # ── About ──
        doc.add_heading("About the Developer", level=1)
        doc.add_paragraph(
            "Dr. Silas Joy Agbesi — Data Scientist Intern, Medical Officer, "
            "Deputy Clinical Coordinator. Dr. Agbesi bridges clinical practice and "
            "data science to build tools that are clinically grounded, explainable, "
            "and safe for real-world healthcare applications."
        )

        # ── Appendix ──
        doc.add_heading("Appendix: Parameters", level=1)
        doc.add_paragraph(f"AutoClin Engine Version: 1.0.0")
        doc.add_paragraph(f"Random Seed: 42")
        doc.add_paragraph(f"Detection Mode: {result.selection_mode}")
        doc.add_paragraph(f"Pipeline Duration: {result.duration_ms / 1000:.2f} seconds")

        doc.save(output_path)
        return output_path
